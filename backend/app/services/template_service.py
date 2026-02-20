import uuid
from pathlib import Path

from docx import Document
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.models.extraction_template import ExtractionTemplate


async def upload_template(
    db: AsyncSession,
    file_bytes: bytes,
    filename: str,
    name: str,
    user_id: uuid.UUID,
    description: str | None = None,
) -> ExtractionTemplate:
    file_id = str(uuid.uuid4())
    file_path = settings.upload_path / f"templates/{file_id}.docx"
    file_path.parent.mkdir(parents=True, exist_ok=True)
    file_path.write_bytes(file_bytes)

    parsed_schema = parse_word_template(str(file_path))

    template = ExtractionTemplate(
        name=name,
        description=description,
        uploaded_by=user_id,
        file_path=str(file_path),
        parsed_schema=parsed_schema,
    )
    db.add(template)
    await db.flush()
    return template


def parse_word_template(file_path: str) -> dict:
    """Parse a Word document to extract the extraction schema.

    The parser identifies:
    - Section headings → extraction categories
    - Tables with columns → fields within categories
    - Placeholder text in cells → field descriptions/instructions
    """
    doc = Document(file_path)
    schema = {"sections": [], "tables": []}

    current_section = None

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            # Check for headings
            for para in doc.paragraphs:
                if para._element is element:
                    if para.style and para.style.name.startswith("Heading"):
                        level = int(para.style.name.replace("Heading ", "").replace("Heading", "1"))
                        current_section = {
                            "name": para.text.strip(),
                            "level": level,
                            "fields": [],
                        }
                        schema["sections"].append(current_section)
                    elif current_section and para.text.strip():
                        # Body text under a heading — treat as field description
                        current_section["fields"].append({
                            "name": para.text.strip(),
                            "type": "text",
                            "description": "",
                        })
                    break

    # Parse tables
    for table_idx, table in enumerate(doc.tables):
        table_data = {"index": table_idx, "columns": [], "rows": []}

        # First row as headers
        if table.rows:
            header_row = table.rows[0]
            table_data["columns"] = [
                cell.text.strip() for cell in header_row.cells
            ]

            # Remaining rows as data/placeholders
            for row in table.rows[1:]:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data["rows"].append(row_data)

        schema["tables"].append(table_data)

        # Add table columns as fields to the current section
        if current_section:
            for col_name in table_data["columns"]:
                if col_name:
                    current_section["fields"].append({
                        "name": col_name,
                        "type": "table_column",
                        "description": f"Column from table {table_idx + 1}",
                    })

    return schema


async def get_templates(db: AsyncSession) -> list[ExtractionTemplate]:
    result = await db.execute(
        select(ExtractionTemplate).order_by(ExtractionTemplate.created_at.desc())
    )
    return list(result.scalars().all())


async def get_template(db: AsyncSession, template_id: uuid.UUID) -> ExtractionTemplate | None:
    result = await db.execute(
        select(ExtractionTemplate).where(ExtractionTemplate.id == template_id)
    )
    return result.scalar_one_or_none()
