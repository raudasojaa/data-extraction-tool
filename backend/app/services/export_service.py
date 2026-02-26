import uuid
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.config import settings
from app.models.article import Article
from app.models.extraction import Extraction
from app.models.grade_assessment import GradeAssessment


CERTAINTY_LABELS = {
    "high": "HIGH",
    "moderate": "MODERATE",
    "low": "LOW",
    "very_low": "VERY LOW",
}

CONFIDENCE_MARKERS = {
    "high": "",
    "medium": " [medium confidence]",
    "low": " [LOW CONFIDENCE]",
}

MISSING_LABELS = {
    "not_reported": "Not reported",
    "explicitly_absent": "Explicitly absent",
    "not_applicable": "N/A",
    "unclear": "Unclear",
}


async def export_extraction_to_word(
    db: AsyncSession,
    extraction_id: uuid.UUID,
) -> Path:
    """Export a single extraction with GRADE assessment to a Word document."""
    result = await db.execute(
        select(Extraction)
        .options(selectinload(Extraction.grade_assessments))
        .where(Extraction.id == extraction_id)
    )
    extraction = result.scalar_one_or_none()
    if not extraction:
        raise ValueError("Extraction not found")

    article_result = await db.execute(
        select(Article).where(Article.id == extraction.article_id)
    )
    article = article_result.scalar_one()

    doc = Document()
    _build_extraction_document(doc, article, extraction)

    output_path = settings.export_path / f"{uuid.uuid4()}.docx"
    doc.save(str(output_path))
    return output_path


async def export_project_to_word(
    db: AsyncSession,
    project_id: uuid.UUID,
) -> Path:
    """Export all extractions in a project to a single Word document."""
    articles_result = await db.execute(
        select(Article)
        .where(Article.project_id == project_id)
        .order_by(Article.created_at)
    )
    articles = list(articles_result.scalars().all())

    doc = Document()
    doc.add_heading("Evidence Synthesis Report", level=0)

    for article in articles:
        ext_result = await db.execute(
            select(Extraction)
            .options(selectinload(Extraction.grade_assessments))
            .where(Extraction.article_id == article.id)
            .order_by(Extraction.version.desc())
            .limit(1)
        )
        extraction = ext_result.scalar_one_or_none()
        if extraction:
            _build_extraction_document(doc, article, extraction)
            doc.add_page_break()

    output_path = settings.export_path / f"project_{uuid.uuid4()}.docx"
    doc.save(str(output_path))
    return output_path


def _build_extraction_document(
    doc: Document,
    article: Article,
    extraction: Extraction,
) -> None:
    """Build the extraction content in a Word document."""
    # Article header
    doc.add_heading(article.title or "Untitled Article", level=1)
    if article.authors:
        doc.add_paragraph(f"Authors: {article.authors}")
    if article.journal:
        doc.add_paragraph(f"Journal: {article.journal} ({article.year or 'N/A'})")

    # Completeness summary table
    if extraction.completeness_summary:
        _build_completeness_table(doc, extraction.completeness_summary)

    # Validation warnings
    if extraction.validation_warnings:
        _build_validation_warnings(doc, extraction.validation_warnings)

    # Study design
    if extraction.study_design:
        doc.add_heading("Study Design", level=2)
        _add_field_data(doc, extraction.study_design)

    # Population
    if extraction.population:
        doc.add_heading("Population", level=2)
        _add_field_data(doc, extraction.population)

    # Intervention
    if extraction.intervention:
        doc.add_heading("Intervention", level=2)
        _add_field_data(doc, extraction.intervention)

    # Comparator
    if extraction.comparator:
        doc.add_heading("Comparator", level=2)
        _add_field_data(doc, extraction.comparator)

    # Outcomes
    if extraction.outcomes:
        doc.add_heading("Outcomes", level=2)
        outcomes = extraction.outcomes
        if isinstance(outcomes, list):
            for outcome in outcomes:
                _add_field_data(doc, outcome)
                doc.add_paragraph("")
        else:
            _add_field_data(doc, outcomes)

    # Additional fields
    for field_name in ["setting", "follow_up", "funding", "limitations", "conclusions"]:
        field_data = getattr(extraction, field_name, None)
        if field_data:
            doc.add_heading(field_name.replace("_", " ").title(), level=2)
            _add_field_data(doc, field_data)

    # Custom fields from template
    if extraction.custom_fields:
        doc.add_heading("Additional Extracted Data", level=2)
        _add_field_data(doc, extraction.custom_fields)

    # GRADE Assessment table
    if extraction.grade_assessments:
        doc.add_heading("GRADE Evidence Profile", level=2)
        _build_grade_table(doc, extraction.grade_assessments)

    # Synthesis
    if extraction.synthesis:
        _build_synthesis_section(doc, extraction.synthesis)


def _add_field_data(doc: Document, data: dict | list) -> None:
    """Add extracted field data to the document, with confidence markers."""
    if isinstance(data, list):
        for item in data:
            if isinstance(item, dict):
                _add_field_data(doc, item)
            else:
                doc.add_paragraph(str(item), style="List Bullet")
        return

    for key, value in data.items():
        if key in ("source_locations", "quotes"):
            continue

        # Handle new {value, confidence, missing_reason, quotes} format
        if isinstance(value, dict) and "confidence" in value:
            _add_confidence_field(doc, key, value)
        elif isinstance(value, dict):
            doc.add_paragraph(f"{key.replace('_', ' ').title()}:")
            _add_field_data(doc, value)
        elif isinstance(value, list):
            doc.add_paragraph(f"{key.replace('_', ' ').title()}:")
            for item in value:
                if isinstance(item, dict):
                    _add_field_data(doc, item)
                else:
                    doc.add_paragraph(f"  - {item}", style="List Bullet")
        else:
            doc.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")


def _add_confidence_field(doc: Document, key: str, field: dict) -> None:
    """Add a field with confidence marker and missing reason styling."""
    label = key.replace("_", " ").title()
    confidence = field.get("confidence")
    missing_reason = field.get("missing_reason")
    value = field.get("value")

    para = doc.add_paragraph()

    # Label run (bold)
    label_run = para.add_run(f"{label}: ")
    label_run.bold = True

    if missing_reason and not value:
        # Missing field
        missing_label = MISSING_LABELS.get(missing_reason, missing_reason)
        value_run = para.add_run(f"[{missing_label}]")
        value_run.italic = True
        value_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    else:
        # Value with confidence marker
        value_run = para.add_run(str(value) if value is not None else "N/A")

        if confidence and confidence != "high":
            marker = CONFIDENCE_MARKERS.get(confidence, "")
            conf_run = para.add_run(marker)
            if confidence == "low":
                conf_run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
                conf_run.bold = True
            elif confidence == "medium":
                conf_run.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)


def _build_completeness_table(doc: Document, summary: dict) -> None:
    """Build an extraction completeness summary table."""
    doc.add_heading("Extraction Completeness", level=2)

    total = summary.get("total_fields", 0)
    extracted = summary.get("extracted", 0)
    missing = summary.get("missing", 0)
    high = summary.get("high_confidence", 0)
    medium = summary.get("medium_confidence", 0)
    low = summary.get("low_confidence", 0)

    table = doc.add_table(rows=2, cols=6)
    table.style = "Table Grid"

    headers = ["Total Fields", "Extracted", "Missing", "High Conf.", "Medium Conf.", "Low Conf."]
    values = [str(total), str(extracted), str(missing), str(high), str(medium), str(low)]

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)

    for i, val in enumerate(values):
        table.rows[1].cells[i].text = val

    # Missing reasons breakdown
    missing_reasons = summary.get("missing_reasons", {})
    if any(v > 0 for v in missing_reasons.values()):
        para = doc.add_paragraph()
        para.add_run("Missing data breakdown: ").bold = True
        parts = []
        for reason, count in missing_reasons.items():
            if count > 0:
                parts.append(f"{MISSING_LABELS.get(reason, reason)}: {count}")
        para.add_run(", ".join(parts))


def _build_validation_warnings(doc: Document, warnings: list) -> None:
    """Add validation warnings to the document."""
    if not warnings:
        return

    doc.add_heading("Validation Warnings", level=3)
    for warning in warnings:
        para = doc.add_paragraph(style="List Bullet")
        severity = warning.get("severity", "warning")
        message = warning.get("message", "")
        field_path = warning.get("field_path", "")

        if severity == "error":
            run = para.add_run(f"ERROR: ")
            run.bold = True
            run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
        else:
            run = para.add_run(f"Warning: ")
            run.bold = True
            run.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)

        para.add_run(f"{message}")
        if field_path:
            field_run = para.add_run(f" ({field_path})")
            field_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def _build_synthesis_section(doc: Document, synthesis: dict) -> None:
    """Add evidence synthesis section to the document."""
    doc.add_heading("Evidence Synthesis", level=2)

    sections = [
        ("key_findings", "Key Findings"),
        ("certainty_of_evidence", "Certainty of Evidence"),
        ("strengths", "Strengths"),
        ("limitations", "Limitations"),
        ("clinical_implications", "Clinical Implications"),
    ]

    for key, label in sections:
        content = synthesis.get(key)
        if content:
            doc.add_heading(label, level=3)
            doc.add_paragraph(content)


def _build_grade_table(doc: Document, assessments: list[GradeAssessment]) -> None:
    """Build a GRADE evidence profile table."""
    headers = [
        "Outcome",
        "Risk of Bias",
        "Inconsistency",
        "Indirectness",
        "Imprecision",
        "Publication Bias",
        "Overall Certainty",
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for assessment in assessments:
        row = table.add_row()
        row.cells[0].text = assessment.outcome_name

        domains = [
            assessment.risk_of_bias,
            assessment.inconsistency,
            assessment.indirectness,
            assessment.imprecision,
            assessment.publication_bias,
        ]
        for i, domain in enumerate(domains):
            if domain and isinstance(domain, dict):
                rating = domain.get("rating", "N/A")
                row.cells[i + 1].text = rating.replace("_", " ").title()

        certainty = assessment.overall_certainty or "N/A"
        row.cells[6].text = CERTAINTY_LABELS.get(certainty, certainty.upper())
