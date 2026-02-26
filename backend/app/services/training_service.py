import logging
import uuid

from docx import Document
from sqlalchemy import func, select, text
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.correction import Correction
from app.models.training_example import TrainingExample
from app.models.user import User

logger = logging.getLogger(__name__)


async def _store_embedding(db: AsyncSession, example_id: uuid.UUID, input_text: str) -> None:
    """Compute and store embedding vector for a training example."""
    try:
        from app.ai.example_selector import compute_embedding
        embedding = compute_embedding(input_text)
        await db.execute(
            text("UPDATE training_examples SET embedding_vector = :vec::vector WHERE id = :id"),
            {"vec": str(embedding), "id": example_id},
        )
        await db.flush()
    except Exception as e:
        logger.warning("Failed to store embedding for example %s: %s", example_id, e)


async def create_training_example_from_correction(
    db: AsyncSession,
    correction: Correction,
    extraction_data: dict,
    article_text: str,
) -> TrainingExample | None:
    """Create a training example from a user correction, if the user is a training contributor."""
    # Check if the user is a training contributor
    result = await db.execute(select(User).where(User.id == correction.user_id))
    user = result.scalar_one_or_none()

    if not user or not user.training_contributor:
        return None

    # Build the corrected output by merging the extraction with the correction
    corrected_output = dict(extraction_data)
    _apply_correction_to_dict(corrected_output, correction.field_path, correction.corrected_value)

    truncated_text = article_text[:10000]
    example = TrainingExample(
        source_type="corrected_extraction",
        source_id=correction.id,
        contributed_by=correction.user_id,
        input_text=truncated_text,
        expected_output=corrected_output,
        quality_score=1.0,
    )
    db.add(example)
    await db.flush()

    # Compute and store embedding vector
    await _store_embedding(db, example.id, truncated_text)

    return example


def _apply_correction_to_dict(data: dict, field_path: str, value) -> None:
    """Apply a correction to a nested dict using dot-notation path."""
    parts = field_path.split(".")
    current = data
    for part in parts[:-1]:
        if "[" in part:
            key, idx = part.split("[")
            idx = int(idx.rstrip("]"))
            current = current.setdefault(key, [])[idx]
        else:
            current = current.setdefault(part, {})

    last = parts[-1]
    if "[" in last:
        key, idx = last.split("[")
        idx = int(idx.rstrip("]"))
        current.setdefault(key, [])[idx] = value
    else:
        current[last] = value


async def import_word_doc_as_training(
    db: AsyncSession,
    file_bytes: bytes,
    user_id: uuid.UUID,
) -> list[TrainingExample]:
    """Parse a completed GRADE assessment Word document into training examples."""
    import io

    doc = Document(io.BytesIO(file_bytes))
    examples = []

    for table in doc.tables:
        parsed = _parse_grade_table(table)
        if parsed:
            example = TrainingExample(
                source_type="imported_word_doc",
                contributed_by=user_id,
                input_text=parsed.get("context", ""),
                expected_output=parsed,
                study_type=parsed.get("study_design"),
                quality_score=0.8,
            )
            db.add(example)
            examples.append(example)

    # Also extract non-table content as narrative training examples
    narrative = _extract_narrative_content(doc)
    if narrative:
        example = TrainingExample(
            source_type="imported_word_doc",
            contributed_by=user_id,
            input_text=narrative[:10000],
            expected_output={"narrative_synthesis": narrative},
            quality_score=0.7,
        )
        db.add(example)
        examples.append(example)

    await db.flush()

    # Compute embeddings for all new examples
    for example in examples:
        if example.input_text:
            await _store_embedding(db, example.id, example.input_text)

    return examples


def _parse_grade_table(table) -> dict | None:
    """Parse a GRADE evidence table from a Word document table."""
    if len(table.rows) < 2 or len(table.columns) < 3:
        return None

    headers = [cell.text.strip().lower() for cell in table.rows[0].cells]

    # Check if this looks like a GRADE table
    grade_keywords = ["outcome", "risk", "certainty", "quality", "evidence", "grade"]
    if not any(kw in " ".join(headers) for kw in grade_keywords):
        return None

    parsed = {"type": "grade_table", "headers": headers, "outcomes": []}

    for row in table.rows[1:]:
        row_data = {}
        for i, cell in enumerate(row.cells):
            if i < len(headers):
                row_data[headers[i]] = cell.text.strip()
        if any(v for v in row_data.values()):
            parsed["outcomes"].append(row_data)

    parsed["context"] = str(parsed)
    return parsed if parsed["outcomes"] else None


def _extract_narrative_content(doc: Document) -> str:
    """Extract narrative/paragraph content from a Word document."""
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and len(text) > 20:
            parts.append(text)
    return "\n\n".join(parts)


async def get_training_stats(db: AsyncSession) -> dict:
    """Get training data statistics."""
    total = await db.scalar(select(func.count(TrainingExample.id)))
    active = await db.scalar(
        select(func.count(TrainingExample.id)).where(TrainingExample.is_active.is_(True))
    )
    avg_quality = await db.scalar(
        select(func.avg(TrainingExample.quality_score)).where(
            TrainingExample.is_active.is_(True)
        )
    )

    # Count by source type
    source_type_result = await db.execute(
        select(TrainingExample.source_type, func.count(TrainingExample.id))
        .group_by(TrainingExample.source_type)
    )
    by_source = dict(source_type_result.all())

    # Count by study type
    study_type_result = await db.execute(
        select(TrainingExample.study_type, func.count(TrainingExample.id))
        .where(TrainingExample.study_type.isnot(None))
        .group_by(TrainingExample.study_type)
    )
    by_study = dict(study_type_result.all())

    return {
        "total_examples": total or 0,
        "active_examples": active or 0,
        "by_source_type": by_source,
        "by_study_type": by_study,
        "avg_quality_score": round(avg_quality or 0, 3),
    }
